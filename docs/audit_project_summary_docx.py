from pathlib import Path
import sys
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


DOCX = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else Path(__file__).resolve().with_name("mall-swarm-project-improvement-summary-and-next-steps.docx")

with ZipFile(DOCX) as archive:
    bad_entry = archive.testzip()
    names = archive.namelist()

doc = Document(DOCX)
text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

print("zip_ok=", bad_entry is None)
print("has_document_xml=", "word/document.xml" in names)
for needle in ["项目改造目标", "已完成的主要改进", "推荐算法评估与看板展示", "后续任务清单", "答辩可强调的创新点"]:
    print(f"{needle}=", needle in text)
print("paragraphs=", len(doc.paragraphs))
print("tables=", len(doc.tables))
print("empty_tables=", sum(1 for table in doc.tables if not table.rows))
print("todo_placeholders=", any(token in text for token in ["TODO", "TBD", "待补充"]))

tables_missing_grid = []
for index, table in enumerate(doc.tables, start=1):
    grid = table._tbl.tblGrid
    grid_cols = [] if grid is None else [column.get(qn("w:w")) for column in grid.iterchildren()]
    if not grid_cols:
        tables_missing_grid.append(index)
print("tables_missing_grid=", tables_missing_grid)
