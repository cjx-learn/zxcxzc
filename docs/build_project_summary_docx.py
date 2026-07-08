from datetime import date
from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "project-improvement-summary-and-next-steps.md"
OUTPUT = ROOT / "docs" / "mall-swarm-project-improvement-summary-and-next-steps.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(102, 112, 133)
BLACK = RGBColor(23, 32, 51)
GRAY_FILL = "F2F4F7"
CALLOUT_FILL = "F8FAFC"
BORDER = "D9E2EE"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("mall-swarm 项目改造说明")
    set_run_font(run, size=9, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer.runs[-1]._r.append(fld_begin)
    footer.runs[-1]._r.append(instr)
    footer.runs[-1]._r.append(fld_end)
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_masthead(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4, line=1.0)
    run = p.add_run("项目改造说明与后续计划")
    set_run_font(run, size=24, color=NAVY, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14, line=1.15)
    run = p.add_run("基于 mall-swarm 的电商用户行为分析与商品推荐平台")
    set_run_font(run, size=13, color=MUTED)

    rows = [
        ("项目", "mall-swarm 二次开发课程作业"),
        ("主题", "基于电商用户行为数据的商品推荐与用户行为分析系统设计与实现"),
        ("版本", "阶段性改造总结"),
        ("日期", "2026-07-07"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1600, 7760])
    for row_idx, (label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        set_cell_shading(cells[0], GRAY_FILL)
        for idx, text in enumerate((label, value)):
            paragraph = cells[idx].paragraphs[0]
            paragraph.text = ""
            run = paragraph.add_run(text)
            set_run_font(run, size=10.5, color=BLACK, bold=(idx == 0))
            set_paragraph_spacing(paragraph, after=0, line=1.10)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=8, line=1.10)
    run = p.add_run(
        "本文件汇总当前已经完成的项目改造、部署状态、核心功能、遗留问题和后续任务，"
        "可作为课程报告、答辩分工和后续开发计划的基础材料。"
    )
    set_run_font(run, size=11, color=BLACK)


def parse_inline(text):
    return re.sub(r"`([^`]+)`", r"\1", text).strip()


def add_paragraph_with_code_runs(paragraph, text, size=11, color=BLACK, bold=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, color=DARK_BLUE, bold=bold, name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    set_paragraph_spacing(paragraph, after=0, line=1.10)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=9.5, color=DARK_BLUE, name="Consolas")
    doc.add_paragraph()


def split_table_row(line):
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [parse_inline(cell) for cell in cells]


def column_widths(rows):
    cols = len(rows[0])
    max_lengths = []
    for col in range(cols):
        length = max(len(str(row[col])) for row in rows if col < len(row))
        max_lengths.append(max(length, 4))
    total = sum(max_lengths)
    raw = [max(1100, int(9360 * length / total)) for length in max_lengths]
    if cols == 2:
        raw = [2600, 6760]
    elif cols == 3:
        raw = [1800, 3000, 4560]
    elif cols == 4:
        raw = [1200, 2200, 2400, 3560]
    delta = 9360 - sum(raw)
    raw[-1] += delta
    return raw


def add_markdown_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = column_widths(rows)
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, GRAY_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            set_paragraph_spacing(paragraph, after=0, line=1.08)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.2 if len(row) >= 4 else 9.8, color=BLACK, bold=(row_idx == 0))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 or row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


def add_body_from_markdown(doc, markdown):
    lines = markdown.splitlines()
    i = 0
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            set_paragraph_spacing(p)
            add_paragraph_with_code_runs(p, " ".join(paragraph_buffer))
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            flush_paragraph()
            i += 1
            continue

        if line.startswith("## "):
            flush_paragraph()
            p = doc.add_paragraph(line[3:].strip(), style="Heading 1")
            set_paragraph_spacing(p, before=16, after=8)
            i += 1
            continue

        if line.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(line[4:].strip(), style="Heading 2")
            set_paragraph_spacing(p, before=12, after=6)
            i += 1
            continue

        if line.startswith("```"):
            flush_paragraph()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            flush_paragraph()
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows)
            continue

        if re.match(r"^\d+\. ", line):
            flush_paragraph()
            item = re.sub(r"^\d+\. ", "", line).strip()
            p = doc.add_paragraph(style="List Number")
            add_paragraph_with_code_runs(p, item)
            set_paragraph_spacing(p, after=4, line=1.167)
            i += 1
            continue

        if line.startswith("- "):
            flush_paragraph()
            item = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_paragraph_with_code_runs(p, item)
            set_paragraph_spacing(p, after=4, line=1.167)
            i += 1
            continue

        paragraph_buffer.append(line.strip())
        i += 1

    flush_paragraph()


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    output = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else OUTPUT
    doc = Document()
    style_document(doc)
    add_masthead(doc)
    add_body_from_markdown(doc, markdown)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
